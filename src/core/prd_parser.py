"""
PRDParser - PRD 文档解析器

支持解析多种格式的 PRD (产品需求文档):
- Markdown (.md)
- Word 文档 (.docx)
- 纯文本 (.txt)

解析后返回纯文本内容，供 LLM 进行业务场景和规则提取。
"""

import logging
import re
from pathlib import Path
from typing import Optional, Dict, Any, List

logger = logging.getLogger(__name__)


class PRDParser:
    """PRD 文档解析器

    将不同格式的 PRD 文档统一解析为纯文本内容。
    """

    # 支持的文件格式
    SUPPORTED_EXTENSIONS = {'.md', '.markdown', '.docx', '.txt', '.text'}

    @staticmethod
    def parse(file_path: str) -> str:
        """解析 PRD 文档

        Args:
            file_path: 文档文件路径

        Returns:
            解析后的纯文本内容

        Raises:
            ValueError: 不支持的文件格式
            FileNotFoundError: 文件不存在
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"PRD 文档不存在: {file_path}")

        ext = path.suffix.lower()
        if ext not in PRDParser.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"不支持的 PRD 格式: {ext}，"
                f"仅支持: {', '.join(PRDParser.SUPPORTED_EXTENSIONS)}"
            )

        logger.info(f"解析 PRD 文档: {file_path} (格式: {ext})")

        if ext in ('.md', '.markdown'):
            return PRDParser._parse_markdown(file_path)
        elif ext == '.docx':
            return PRDParser._parse_docx(file_path)
        elif ext in ('.txt', '.text'):
            return PRDParser._parse_txt(file_path)
        else:
            return PRDParser._parse_txt(file_path)

    @staticmethod
    def _parse_markdown(file_path: str) -> str:
        """解析 Markdown 文件

        保留 Markdown 结构，因为 LLM 可以很好地理解 Markdown 格式。

        Args:
            file_path: Markdown 文件路径

        Returns:
            Markdown 原文内容
        """
        encodings = ['utf-8', 'utf-8-sig', 'gbk', 'gb2312', 'latin-1']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                    logger.info(f"  Markdown 解析成功，长度: {len(content)} 字符")
                    return content
            except UnicodeDecodeError:
                continue

        raise ValueError(f"无法解析 Markdown 文件，尝试的编码: {encodings}")

    @staticmethod
    def _parse_docx(file_path: str) -> str:
        """解析 Word 文档 (.docx)

        提取所有段落文本，保留基本结构。

        Args:
            file_path: Word 文档路径

        Returns:
            提取的纯文本内容
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("需要安装 python-docx: pip install python-docx")

        doc = Document(file_path)
        lines = []

        # 提取所有段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # 根据样式判断是否为标题
                style_name = para.style.name.lower() if para.style else ""
                if 'heading' in style_name or 'title' in style_name:
                    # 转换为 Markdown 标题格式
                    level = 1
                    if 'heading 1' in style_name:
                        level = 1
                    elif 'heading 2' in style_name:
                        level = 2
                    elif 'heading 3' in style_name:
                        level = 3
                    elif 'heading' in style_name:
                        level = 2
                    lines.append(f"\n{'#' * level} {text}\n")
                else:
                    lines.append(text)

        # 提取表格内容
        for table in doc.tables:
            lines.append("\n[表格内容]")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    lines.append(" | ".join(cells))
            lines.append("")

        content = "\n".join(lines)
        logger.info(f"  Word 文档解析成功，长度: {len(content)} 字符")
        return content

    @staticmethod
    def _parse_txt(file_path: str) -> str:
        """解析纯文本文件

        Args:
            file_path: 文本文件路径

        Returns:
            文本内容
        """
        encodings = ['utf-8', 'utf-8-sig', 'gbk', 'gb2312', 'latin-1']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                    logger.info(f"  文本文件解析成功，长度: {len(content)} 字符")
                    return content
            except UnicodeDecodeError:
                continue

        raise ValueError(f"无法解析文本文件，尝试的编码: {encodings}")

    @staticmethod
    def extract_sections(content: str) -> Dict[str, str]:
        """提取文档章节

        将 PRD 文档按标题分割为章节，便于后续处理。

        Args:
            content: 文档内容

        Returns:
            Dict[str, str]: {章节标题: 章节内容}
        """
        sections = {}

        # 匹配 Markdown 标题
        pattern = r'^(#{1,3})\s+(.+)$'
        lines = content.split('\n')

        current_title = "概述"
        current_content = []

        for line in lines:
            match = re.match(pattern, line)
            if match:
                # 保存上一个章节
                if current_content:
                    sections[current_title] = '\n'.join(current_content).strip()

                # 开始新章节
                current_title = match.group(2).strip()
                current_content = []
            else:
                current_content.append(line)

        # 保存最后一个章节
        if current_content:
            sections[current_title] = '\n'.join(current_content).strip()

        return sections

    @staticmethod
    def summarize_for_prompt(content: str, max_length: int = 10000) -> str:
        """生成 PRD 摘要供 LLM Prompt 使用

        如果内容过长，进行智能截断，保留关键部分。

        Args:
            content: PRD 原文内容
            max_length: 最大长度限制

        Returns:
            处理后的内容
        """
        if len(content) <= max_length:
            return content

        logger.info(f"PRD 内容过长 ({len(content)} 字符)，进行智能截断")

        # 提取章节
        sections = PRDParser.extract_sections(content)

        # 优先保留的关键词
        priority_keywords = [
            '流程', '规则', '业务', '功能', '接口', '场景',
            'flow', 'rule', 'business', 'function', 'api', 'scenario'
        ]

        # 按优先级排序章节
        prioritized = []
        other = []

        for title, text in sections.items():
            title_lower = title.lower()
            if any(kw in title_lower for kw in priority_keywords):
                prioritized.append((title, text))
            else:
                other.append((title, text))

        # 构建摘要
        result_parts = []
        remaining_length = max_length

        # 先添加优先章节
        for title, text in prioritized:
            section_text = f"## {title}\n{text}\n"
            if len(section_text) <= remaining_length:
                result_parts.append(section_text)
                remaining_length -= len(section_text)
            elif remaining_length > 500:
                # 截断该章节
                truncated = text[:remaining_length - 100] + "\n...(内容已截断)"
                result_parts.append(f"## {title}\n{truncated}\n")
                remaining_length = 0
                break

        # 如果还有空间，添加其他章节
        for title, text in other:
            if remaining_length <= 500:
                break
            section_text = f"## {title}\n{text}\n"
            if len(section_text) <= remaining_length:
                result_parts.append(section_text)
                remaining_length -= len(section_text)

        result = "\n".join(result_parts)
        logger.info(f"PRD 截断后长度: {len(result)} 字符")
        return result

    @staticmethod
    def detect_business_keywords(content: str) -> List[Dict[str, Any]]:
        """检测 PRD 中的业务关键词

        用于初步分析 PRD 中可能包含的业务场景和规则。

        Args:
            content: PRD 内容

        Returns:
            检测到的关键词列表，包含类型和位置
        """
        keywords = []

        # 业务流程关键词
        flow_patterns = [
            (r'流程[：:]\s*(.+)', 'flow'),
            (r'步骤[：:]\s*(.+)', 'step'),
            (r'(\d+)[\.、]\s*用户(.+)', 'user_action'),
            (r'(\d+)[\.、]\s*系统(.+)', 'system_action'),
        ]

        # 业务规则关键词
        rule_patterns = [
            (r'规则[：:]\s*(.+)', 'rule'),
            (r'如果(.+)[,，]则(.+)', 'condition'),
            (r'当(.+)时[,，](.+)', 'condition'),
            (r'必须(.+)', 'constraint'),
            (r'不能(.+)', 'constraint'),
            (r'(.+)[=＝](.+)', 'formula'),
        ]

        lines = content.split('\n')
        for i, line in enumerate(lines):
            for pattern, keyword_type in flow_patterns + rule_patterns:
                match = re.search(pattern, line)
                if match:
                    keywords.append({
                        'type': keyword_type,
                        'line': i + 1,
                        'text': line.strip(),
                        'match': match.group(0)
                    })

        return keywords

    @staticmethod
    def is_supported(file_path: str) -> bool:
        """检查文件格式是否支持

        Args:
            file_path: 文件路径

        Returns:
            是否支持该格式
        """
        ext = Path(file_path).suffix.lower()
        return ext in PRDParser.SUPPORTED_EXTENSIONS
